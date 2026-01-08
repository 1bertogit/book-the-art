#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Cria um arquivo reference.docx para uso com Pandoc.

Define estilos profissionais para o ebook:
- Headings com fonte sans-serif
- Corpo com fonte serif
- Espaçamento adequado entre parágrafos
- Margens de página configuradas

Uso:
    python3 tools/create_reference_docx.py
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parents[1]
OUTPUT_PATH = ROOT / "assets" / "reference.docx"


def create_reference_doc():
    """Cria o documento de referência com estilos customizados."""
    doc = Document()

    # ===== CONFIGURAR MARGENS =====
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # ===== ESTILO NORMAL (CORPO) =====
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Georgia'
    font.size = Pt(12)  # Aumentado de 11 para 12
    font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)  # Cinza escuro, não preto puro

    para_format = style_normal.paragraph_format
    para_format.space_after = Pt(12)
    para_format.line_spacing = 1.5
    para_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Texto justificado

    # ===== HEADING 1 =====
    style_h1 = doc.styles['Heading 1']
    font = style_h1.font
    font.name = 'Helvetica Neue'
    font.size = Pt(24)
    font.bold = True
    font.color.rgb = RGBColor(0x0f, 0x4c, 0x81)  # Azul escuro

    para_format = style_h1.paragraph_format
    para_format.space_before = Pt(36)
    para_format.space_after = Pt(12)
    para_format.page_break_before = True  # Capítulo sempre em nova página

    # ===== HEADING 2 =====
    style_h2 = doc.styles['Heading 2']
    font = style_h2.font
    font.name = 'Helvetica Neue'
    font.size = Pt(18)
    font.bold = True
    font.color.rgb = RGBColor(0x2e, 0x74, 0xb5)  # Azul médio

    para_format = style_h2.paragraph_format
    para_format.space_before = Pt(24)
    para_format.space_after = Pt(8)

    # ===== HEADING 3 =====
    style_h3 = doc.styles['Heading 3']
    font = style_h3.font
    font.name = 'Helvetica Neue'
    font.size = Pt(14)
    font.bold = True
    font.color.rgb = RGBColor(0x4a, 0x4a, 0x4a)  # Cinza escuro

    para_format = style_h3.paragraph_format
    para_format.space_before = Pt(18)
    para_format.space_after = Pt(6)

    # ===== BLOCKQUOTE (Quote) =====
    try:
        style_quote = doc.styles['Quote']
    except KeyError:
        style_quote = doc.styles.add_style('Quote', WD_STYLE_TYPE.PARAGRAPH)

    font = style_quote.font
    font.name = 'Georgia'
    font.size = Pt(10)
    font.italic = True
    font.color.rgb = RGBColor(0x55, 0x55, 0x55)  # Cinza

    para_format = style_quote.paragraph_format
    para_format.left_indent = Inches(0.5)
    para_format.space_before = Pt(12)
    para_format.space_after = Pt(12)

    # ===== LIST BULLET =====
    try:
        style_list = doc.styles['List Bullet']
        font = style_list.font
        font.name = 'Georgia'
        font.size = Pt(11)

        para_format = style_list.paragraph_format
        para_format.space_after = Pt(6)
        para_format.left_indent = Inches(0.5)
    except KeyError:
        pass

    # ===== CAPTION (para figuras) =====
    try:
        style_caption = doc.styles['Caption']
        font = style_caption.font
        font.name = 'Helvetica Neue'
        font.size = Pt(10)
        font.italic = True
        font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        para_format = style_caption.paragraph_format
        para_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para_format.space_before = Pt(6)
        para_format.space_after = Pt(12)
    except KeyError:
        pass

    # ===== TITLE =====
    try:
        style_title = doc.styles['Title']
        font = style_title.font
        font.name = 'Helvetica Neue'
        font.size = Pt(36)
        font.bold = True
        font.color.rgb = RGBColor(0x0f, 0x4c, 0x81)

        para_format = style_title.paragraph_format
        para_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para_format.space_after = Pt(24)
    except KeyError:
        pass

    # ===== SUBTITLE =====
    try:
        style_subtitle = doc.styles['Subtitle']
        font = style_subtitle.font
        font.name = 'Georgia'
        font.size = Pt(18)
        font.italic = True
        font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        para_format = style_subtitle.paragraph_format
        para_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para_format.space_after = Pt(36)
    except KeyError:
        pass

    # Adicionar parágrafos de exemplo para cada estilo
    # (Pandoc precisa ver os estilos em uso)
    doc.add_heading('Título Exemplo', level=1)
    doc.add_heading('Seção Exemplo', level=2)
    doc.add_heading('Subseção Exemplo', level=3)
    doc.add_paragraph('Parágrafo de texto normal para definir o estilo base.')

    # Salvar
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_PATH))

    print(f"✅ Reference DOCX criado: {OUTPUT_PATH}")
    print()
    print("Estilos configurados:")
    print("  • Heading 1: Helvetica Neue 24pt, azul escuro")
    print("  • Heading 2: Helvetica Neue 18pt, azul médio")
    print("  • Heading 3: Helvetica Neue 14pt, cinza")
    print("  • Normal: Georgia 12pt, espaçamento 1.5, justificado")
    print("  • Quote: Georgia 10pt italic, indentado")
    print("  • Margens: 1 polegada (2.54cm)")


if __name__ == "__main__":
    create_reference_doc()
