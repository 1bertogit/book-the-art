#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Builder de eBook com Python a partir de Markdown.
Suporta DOCX, EPUB e HTML.

Uso:
  python build.py --format docx --out dist/ebook.docx
  python build.py --format epub --out dist/ebook.epub
  python build.py --format html --out dist/ebook.html

Dica (recomendado): instale o Pandoc para conversão de alta fidelidade.
"""

import os
import sys
import argparse
import glob
import yaml

import markdown as md  # HTML simples (fallback)

try:
    import pypandoc  # Conversão superior (DOCX/EPUB/HTML)
    HAS_PANDOC = True
except Exception:
    HAS_PANDOC = False

try:
    from ebooklib import epub
    HAS_EBOOKLIB = True
except Exception:
    HAS_EBOOKLIB = False

try:
    from docx import Document
    HAS_DOCPY = True
except Exception:
    HAS_DOCPY = False


def read_config(path="config.yml"):
    if not os.path.exists(path):
        return {}
    with open(path, "r", encoding="utf-8") as f:
        return yaml.safe_load(f) or {}


def collect_markdown(md_dir="content"):
    files = sorted(glob.glob(os.path.join(md_dir, "*.md")))
    parts = []
    for fp in files:
        with open(fp, "r", encoding="utf-8") as f:
            parts.append(f.read())
    return parts, files


def ensure_dist(path):
    dist = os.path.dirname(path)
    if dist and not os.path.exists(dist):
        os.makedirs(dist, exist_ok=True)


def build_html(config, md_parts):
    md_text = "\n\n\n".join(md_parts)
    html = md.markdown(md_text, extensions=["toc", "tables", "fenced_code"])
    title = config.get("title", "Meu eBook")
    subtitle = config.get("subtitle", "")
    header = f"<h1>{title}</h1>" + (f"<h2>{subtitle}</h2>" if subtitle else "")
    final = f"""<!DOCTYPE html>
<html lang="{config.get('language','pt-BR')}">
<head>
<meta charset="utf-8" />
<title>{title}</title>
<meta name="viewport" content="width=device-width, initial-scale=1" />
<style>
body {{ font-family: -apple-system, Segoe UI, Roboto, Helvetica, Arial, sans-serif; max-width: 900px; margin: 2rem auto; padding: 0 1rem; line-height: 1.65; }}
h1,h2,h3,h4 {{ line-height: 1.25; }}
pre, code {{ font-family: ui-monospace, SFMono-Regular, Menlo, Consolas, monospace; }}
blockquote {{ border-left: 3px solid #ddd; padding-left: 1rem; color: #444; }}
table {{ border-collapse: collapse; width: 100%; }}
td, th {{ border: 1px solid #ddd; padding: 8px; }}
</style>
</head>
<body>
{header}
{html}
</body>
</html>"""
    return final


def build_docx(config, md_parts, out_path):
    if HAS_PANDOC:
        tmp_md = ".build_tmp.md"
        with open(tmp_md, "w", encoding="utf-8") as f:
            f.write("\n\n\n".join(md_parts))
        pypandoc.convert_file(tmp_md, "docx", outputfile=out_path)
        os.remove(tmp_md)
        return

    if not HAS_DOCPY:
        raise RuntimeError("Para DOCX sem Pandoc, instale python-docx (ou instale Pandoc).")

    # Fallback simplificado com python-docx (Markdown parcial)
    doc = Document()
    doc.add_heading(config.get("title", "Meu eBook"), 0)
    subtitle = config.get("subtitle", "")
    if subtitle:
        doc.add_paragraph(subtitle)

    for part in md_parts:
        for line in part.splitlines():
            if line.startswith("# "):
                doc.add_heading(line[2:].strip(), level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:].strip(), level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:].strip(), level=3)
            else:
                doc.add_paragraph(line)
        doc.add_page_break()

    doc.save(out_path)


def build_epub(config, md_parts, out_path):
    if HAS_PANDOC:
        tmp_md = ".build_tmp.md"
        with open(tmp_md, "w", encoding="utf-8") as f:
            f.write("\n\n\n".join(md_parts))
        pypandoc.convert_file(tmp_md, "epub", outputfile=out_path)
        os.remove(tmp_md)
        return

    if not HAS_EBOOKLIB:
        raise RuntimeError("Para EPUB sem Pandoc, instale ebooklib (ou instale Pandoc).")

    book = epub.EpubBook()
    book.set_identifier("the-art-of-eyelid-surgery")
    book.set_title(config.get("title", "Meu eBook"))
    book.set_language(config.get("language", "pt-BR"))
    book.add_author(config.get("author", "Autor"))

    chapters = []
    for i, part in enumerate(md_parts, start=1):
        html = md.markdown(part, extensions=["toc", "tables", "fenced_code"])
        c = epub.EpubHtml(title=f"Capítulo {i}", file_name=f"chap_{i:02d}.xhtml", lang=config.get("language","pt-BR"))
        c.content = html
        book.add_item(c)
        chapters.append(c)

    book.toc = tuple(chapters)
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())

    style = "BODY { font-family: serif; line-height: 1.4; }"
    nav_css = epub.EpubItem(uid="style_nav", file_name="style/nav.css", media_type="text/css", content=style)
    book.add_item(nav_css)

    book.spine = ["nav"] + chapters
    epub.write_epub(out_path, book)


def main():
    parser = argparse.ArgumentParser(description="Builder de eBook com Python")
    parser.add_argument("--format", choices=["docx", "epub", "html"], required=True)
    parser.add_argument("--out", required=True)
    args = parser.parse_args()

    config = read_config()
    md_parts, files = collect_markdown("content")
    if not md_parts:
        print("Nenhum arquivo .md encontrado em ./content")
        sys.exit(2)

    ensure_dist(args.out)

    if args.format == "html":
        html = build_html(config, md_parts)
        with open(args.out, "w", encoding="utf-8") as f:
            f.write(html)
        print(f"[OK] HTML gerado em: {args.out}")
        return

    if args.format == "docx":
        build_docx(config, md_parts, args.out)
        print(f"[OK] DOCX gerado em: {args.out}")
        return

    if args.format == "epub":
        build_epub(config, md_parts, args.out)
        print(f"[OK] EPUB gerado em: {args.out}")
        return


if __name__ == "__main__":
    main()
